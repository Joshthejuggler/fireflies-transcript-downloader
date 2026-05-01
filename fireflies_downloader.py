#!/usr/bin/env python3
"""
Fireflies Transcript Downloader — GUI
Downloads meeting transcripts as .docx files for a given date range.
Saves state so you always know where you left off.
"""

import subprocess, sys, os, json, threading
from datetime import datetime, timezone, timedelta
from tkinter import filedialog

# ── Auto-install dependencies ─────────────────────────────────────────────────
def pip(pkg):
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q"],
                              stderr=subprocess.DEVNULL)
    except subprocess.CalledProcessError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q",
                               "--break-system-packages"])

for pkg, imp in [("customtkinter","customtkinter"), ("requests","requests"),
                 ("python-docx","docx")]:
    try:
        __import__(imp)
    except ImportError:
        print(f"Installing {pkg}…"); pip(pkg)

import customtkinter as ctk
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Config ────────────────────────────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
STATE_FILE  = os.path.join(SCRIPT_DIR, "fireflies_state.json")
API_URL     = "https://api.fireflies.ai/graphql"
APP_VERSION = "1.0"

DEFAULT_STATE = {
    "api_key": "",          # set on first launch — stored in fireflies_state.json
    "output_dir": SCRIPT_DIR,
    "last_downloaded_date": None,
    "last_downloaded_title": None,
}

# ── State helpers ─────────────────────────────────────────────────────────────
def load_state():
    if os.path.exists(STATE_FILE):
        try:
            s = DEFAULT_STATE.copy()
            with open(STATE_FILE) as f:
                s.update(json.load(f))
            return s
        except Exception:
            pass
    return DEFAULT_STATE.copy()

def save_state(state):
    with open(STATE_FILE, "w") as f:
        json.dump(state, f, indent=2)

# ── Fireflies API ─────────────────────────────────────────────────────────────
def graphql(api_key, query):
    r = requests.post(
        API_URL,
        json={"query": query},
        headers={"Content-Type": "application/json",
                 "Authorization": f"Bearer {api_key}"},
        timeout=30,
    )
    r.raise_for_status()
    data = r.json()
    if "errors" in data:
        raise RuntimeError(data["errors"][0]["message"])
    return data["data"]

def fetch_transcripts(api_key, from_iso, to_iso):
    q = f"""
    {{
      transcripts(fromDate: "{from_iso}", toDate: "{to_iso}", limit: 50) {{
        id title date duration
        sentences {{ speaker_name text }}
      }}
    }}
    """
    return graphql(api_key, q)["transcripts"]

# ── Date / name helpers ───────────────────────────────────────────────────────
def parse_date(val):
    if isinstance(val, (int, float)):
        return datetime.fromtimestamp(val / 1000, tz=timezone.utc)
    return datetime.fromisoformat(str(val).replace("Z", "+00:00"))

def date_prefix(val):  return parse_date(val).strftime("%Y-%m-%d")
def date_long(val):    return parse_date(val).strftime("%B %d, %Y")
def fmt_dur(secs):     return f"{int(secs)//60} min" if secs else ""

def safe_name(s):
    for ch in r'<>:"/\|?*':
        s = s.replace(ch, "")
    return " ".join(s.split())   # collapse extra spaces from stripped chars

def to_iso_start(d): return f"{d}T00:00:00.000Z"
def to_iso_end(d):   return f"{d}T23:59:59.999Z"

def next_day(d):
    """'YYYY-MM-DD' → next day string, capped at today"""
    nxt = datetime.strptime(d, "%Y-%m-%d") + timedelta(days=1)
    today = datetime.now()
    return min(nxt, today).strftime("%Y-%m-%d")

# ── DOCX builder ──────────────────────────────────────────────────────────────
def build_docx(t, path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1.25)

    # Title
    h = doc.add_heading(t["title"], level=1)
    h.runs[0].font.size = Pt(16)
    h.runs[0].font.bold = True
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Date + duration
    meta = doc.add_paragraph()
    info = date_long(t["date"])
    if t.get("duration"):
        info += f"  ·  {fmt_dur(t['duration'])}"
    run = meta.add_run(info)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    meta.paragraph_format.space_after = Pt(14)

    # Horizontal rule
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(0)
    hr.paragraph_format.space_after  = Pt(14)
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    for k, v in [("w:val","single"),("w:sz","6"),("w:space","1"),("w:color","CCCCCC")]:
        bot.set(qn(k), v)
    pBdr.append(bot); pPr.append(pBdr)

    # Transcript — grouped by speaker turn, no timestamps
    sentences = t.get("sentences") or []
    if not sentences:
        doc.add_paragraph("No transcript available.")
    else:
        turns = []
        for s in sentences:
            spk  = (s.get("speaker_name") or "Unknown").strip()
            text = (s.get("text") or "").strip()
            if not text: continue
            if turns and turns[-1][0] == spk:
                turns[-1][1].append(text)
            else:
                turns.append([spk, [text]])

        for spk, texts in turns:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(10)
            sp.paragraph_format.space_after  = Pt(2)
            sr = sp.add_run(spk)
            sr.bold = True
            sr.font.size = Pt(11)
            sr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

            body = doc.add_paragraph(" ".join(texts))
            body.paragraph_format.space_before = Pt(0)
            body.paragraph_format.space_after  = Pt(4)
            body.runs[0].font.size = Pt(11)

    doc.save(path)

# ── GUI ───────────────────────────────────────────────────────────────────────
class App(ctk.CTk):

    def __init__(self):
        super().__init__()
        self.state = load_state()
        ctk.set_appearance_mode("system")
        ctk.set_default_color_theme("blue")
        self.title("Fireflies Downloader")
        self.geometry("700x660")
        self.minsize(600, 560)
        self._build_ui()

    # ── Layout ────────────────────────────────────────────────────────────
    def _build_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(6, weight=1)   # log expands

        # ── Header ───────────────────────────────────────────────────────
        hdr = ctk.CTkFrame(self, corner_radius=0, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=24, pady=(20, 4))
        hdr.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(hdr, text="Fireflies Transcript Downloader",
                     font=ctk.CTkFont(size=22, weight="bold")).grid(
            row=0, column=0, sticky="w")

        self.last_label = ctk.CTkLabel(
            hdr, text=self._last_text(),
            font=ctk.CTkFont(size=12), text_color="gray")
        self.last_label.grid(row=1, column=0, sticky="w", pady=(2, 0))

        # ── Settings card ────────────────────────────────────────────────
        sf = ctk.CTkFrame(self)
        sf.grid(row=1, column=0, sticky="ew", padx=24, pady=(12, 0))
        sf.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(sf, text="API Key", anchor="w",
                     font=ctk.CTkFont(weight="bold")).grid(
            row=0, column=0, sticky="w", padx=14, pady=(12, 4))
        self.api_var = ctk.StringVar(value=self.state["api_key"])
        self.api_entry = ctk.CTkEntry(sf, textvariable=self.api_var, show="•")
        self.api_entry.grid(row=0, column=1, columnspan=2, sticky="ew",
                            padx=(0, 14), pady=(12, 4))

        show_btn = ctk.CTkButton(sf, text="Show", width=56,
                                 fg_color="transparent", border_width=1,
                                 command=self._toggle_key)
        show_btn.grid(row=0, column=3, padx=(0, 14), pady=(12, 4))
        self._key_hidden = True
        self._show_btn = show_btn

        ctk.CTkLabel(sf, text="Save folder", anchor="w",
                     font=ctk.CTkFont(weight="bold")).grid(
            row=1, column=0, sticky="w", padx=14, pady=(4, 14))
        self.out_var = ctk.StringVar(value=self.state["output_dir"])
        ctk.CTkEntry(sf, textvariable=self.out_var).grid(
            row=1, column=1, columnspan=2, sticky="ew", padx=(0, 6), pady=(4, 14))
        ctk.CTkButton(sf, text="Browse", width=70,
                      command=self._browse).grid(
            row=1, column=3, padx=(0, 14), pady=(4, 14))

        # ── Date range card ──────────────────────────────────────────────
        df = ctk.CTkFrame(self)
        df.grid(row=2, column=0, sticky="ew", padx=24, pady=(10, 0))
        df.grid_columnconfigure((1, 3), weight=1)

        ctk.CTkLabel(df, text="Date Range",
                     font=ctk.CTkFont(weight="bold")).grid(
            row=0, column=0, columnspan=4, sticky="w", padx=14, pady=(12, 6))

        ctk.CTkLabel(df, text="From").grid(
            row=1, column=0, sticky="w", padx=14, pady=(0, 12))
        self.from_var = ctk.StringVar(value=self._default_from())
        ctk.CTkEntry(df, textvariable=self.from_var,
                     placeholder_text="YYYY-MM-DD").grid(
            row=1, column=1, sticky="ew", padx=(0, 20), pady=(0, 12))

        ctk.CTkLabel(df, text="To").grid(
            row=1, column=2, sticky="w", padx=(4, 0), pady=(0, 12))
        self.to_var = ctk.StringVar(
            value=datetime.now().strftime("%Y-%m-%d"))
        ctk.CTkEntry(df, textvariable=self.to_var,
                     placeholder_text="YYYY-MM-DD").grid(
            row=1, column=3, sticky="ew", padx=(0, 14), pady=(0, 12))

        # Smart checkbox
        self.use_last_var = ctk.BooleanVar(
            value=bool(self.state.get("last_downloaded_date")))
        ctk.CTkCheckBox(
            df,
            text="Auto-set 'From' to the day after last download",
            variable=self.use_last_var,
            command=self._apply_use_last,
        ).grid(row=2, column=0, columnspan=4, sticky="w", padx=14, pady=(0, 12))

        # ── Download button ──────────────────────────────────────────────
        self.dl_btn = ctk.CTkButton(
            self,
            text="⬇  Download Transcripts",
            font=ctk.CTkFont(size=15, weight="bold"),
            height=48,
            command=self._start_download,
        )
        self.dl_btn.grid(row=3, column=0, sticky="ew", padx=24, pady=(14, 0))

        # ── Progress bar ─────────────────────────────────────────────────
        self.progress = ctk.CTkProgressBar(self)
        self.progress.set(0)
        self.progress.grid(row=4, column=0, sticky="ew", padx=24, pady=(8, 0))

        # ── Status label ─────────────────────────────────────────────────
        self.status_label = ctk.CTkLabel(
            self, text="", font=ctk.CTkFont(size=12), text_color="gray")
        self.status_label.grid(row=5, column=0, sticky="w", padx=24)

        # ── Log ──────────────────────────────────────────────────────────
        self.log = ctk.CTkTextbox(
            self, font=ctk.CTkFont(family="Menlo", size=11), wrap="word")
        self.log.grid(row=6, column=0, sticky="nsew", padx=24, pady=(6, 20))
        self.log.configure(state="disabled")

    # ── Helpers ───────────────────────────────────────────────────────────
    def _last_text(self):
        d = self.state.get("last_downloaded_date")
        t = self.state.get("last_downloaded_title")
        if d and t:
            return f"Last downloaded  ·  {d}  —  {t}"
        return "Last downloaded  ·  none yet"

    def _default_from(self):
        d = self.state.get("last_downloaded_date")
        if d:
            try:
                return next_day(d)
            except Exception:
                pass
        return "2020-01-01"

    def _apply_use_last(self):
        if self.use_last_var.get():
            self.from_var.set(self._default_from())

    def _browse(self):
        folder = filedialog.askdirectory(initialdir=self.out_var.get())
        if folder:
            self.out_var.set(folder)

    def _toggle_key(self):
        self._key_hidden = not self._key_hidden
        self.api_entry.configure(show="•" if self._key_hidden else "")
        self._show_btn.configure(text="Show" if self._key_hidden else "Hide")

    def _log(self, msg, color=None):
        self.log.configure(state="normal")
        self.log.insert("end", msg + "\n")
        self.log.see("end")
        self.log.configure(state="disabled")

    def _set_status(self, msg):
        self.status_label.configure(text=msg)

    # ── Download ──────────────────────────────────────────────────────────
    def _start_download(self):
        self.dl_btn.configure(state="disabled", text="Downloading…")
        self.progress.set(0)
        self.log.configure(state="normal")
        self.log.delete("1.0", "end")
        self.log.configure(state="disabled")
        threading.Thread(target=self._worker, daemon=True).start()

    def _worker(self):
        try:
            api_key    = self.api_var.get().strip()
            output_dir = self.out_var.get().strip()
            from_d     = self.from_var.get().strip()
            to_d       = self.to_var.get().strip()

            # Validate
            if not api_key:
                self._log("✗  API key is required."); return
            if not os.path.isdir(output_dir):
                self._log("✗  Save folder does not exist."); return
            try:
                from_dt = datetime.strptime(from_d, "%Y-%m-%d")
                to_dt   = datetime.strptime(to_d,   "%Y-%m-%d")
            except ValueError:
                self._log("✗  Dates must be in YYYY-MM-DD format."); return
            if from_dt > to_dt:
                self._log(f"✗  'From' date ({from_d}) is after 'To' date ({to_d}).")
                self._log("   Please fix the date range and try again.")
                self._set_status("Invalid date range")
                return

            self._log(f"Connecting to Fireflies API…")
            self._set_status(f"Fetching  {from_d}  →  {to_d}")

            transcripts = fetch_transcripts(
                api_key, to_iso_start(from_d), to_iso_end(to_d))

            if not transcripts:
                self._log("No transcripts found for that date range.")
                self._set_status("Done — nothing to download.")
                return

            total = len(transcripts)
            self._log(f"Found {total} transcript(s)\n")

            newest_date  = None
            newest_title = None
            ok = 0

            for i, t in enumerate(transcripts, 1):
                try:
                    prefix   = date_prefix(t["date"])
                    filename = f"{prefix} {safe_name(t['title'])}.docx"
                    path     = os.path.join(output_dir, filename)

                    self._log(f"  →  {filename}")
                    self._set_status(f"[{i}/{total}]  {filename}")
                    self.progress.set(i / total)

                    build_docx(t, path)
                    ok += 1

                    if newest_date is None or prefix > newest_date:
                        newest_date  = prefix
                        newest_title = t["title"]

                except Exception as e:
                    self._log(f"  ✗  Error on '{t.get('title')}': {e}")

            # Persist state
            if newest_date:
                self.state.update({
                    "api_key": api_key,
                    "output_dir": output_dir,
                    "last_downloaded_date":  newest_date,
                    "last_downloaded_title": newest_title,
                })
                save_state(self.state)
                self.last_label.configure(text=self._last_text())
                self.use_last_var.set(True)
                self._apply_use_last()

            self.progress.set(1)
            self._log(f"\n✓  {ok}/{total} file(s) saved to:\n   {output_dir}")
            self._set_status(f"Done  ·  {ok} file(s) saved")

        except Exception as e:
            self._log(f"\n✗  {e}")
            self._set_status("Error — see log above")
        finally:
            self.dl_btn.configure(
                state="normal", text="⬇  Download Transcripts")


if __name__ == "__main__":
    app = App()
    app.mainloop()
