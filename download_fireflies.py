#!/usr/bin/env python3
"""
Fireflies Transcript Downloader
Downloads meeting transcripts from March 30, 2026 and earlier as .docx files
No timestamps — speaker names + text only
"""

import json
import os
import sys
import subprocess

# ── Auto-install dependencies ────────────────────────────────────────────────
def install(pkg):
    # Try normal install first, then fall back to --break-system-packages (Homebrew Python)
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q"],
                              stderr=subprocess.DEVNULL)
    except subprocess.CalledProcessError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q",
                               "--break-system-packages"])

try:
    import requests
except ImportError:
    print("Installing requests..."); install("requests"); import requests

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx..."); install("python-docx")
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

from datetime import datetime, timezone

# ── Config ───────────────────────────────────────────────────────────────────
API_KEY    = os.environ.get("FIREFLIES_API_KEY", "")  # set via env var or edit here locally
API_URL    = "https://api.fireflies.ai/graphql"
TO_DATE    = "2026-03-31T00:00:00.000Z"   # everything up to & including Mar 30
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── GraphQL helper ───────────────────────────────────────────────────────────
def gql(query, variables=None):
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {API_KEY}",
    }
    payload = {"query": query}
    if variables:
        payload["variables"] = variables
    r = requests.post(API_URL, json=payload, headers=headers, timeout=30)
    r.raise_for_status()
    data = r.json()
    if "errors" in data:
        raise RuntimeError(f"GraphQL error: {data['errors']}")
    return data["data"]

# ── Fetch all transcripts up to Mar 30 ──────────────────────────────────────
def fetch_all_transcripts():
    query = """
    {
      transcripts(toDate: "%s", limit: 50) {
        id title date duration
        sentences { speaker_name text }
      }
    }
    """ % TO_DATE
    return gql(query)["transcripts"]

# ── Date helpers ─────────────────────────────────────────────────────────────
def parse_date(val):
    """Return a datetime from ms-epoch float OR ISO string."""
    if isinstance(val, (int, float)):
        return datetime.fromtimestamp(val / 1000, tz=timezone.utc)
    s = str(val)
    # strip trailing Z for fromisoformat compat on older Pythons
    return datetime.fromisoformat(s.replace("Z", "+00:00"))

def date_prefix(val):
    return parse_date(val).strftime("%Y-%m-%d")

def date_long(val):
    return parse_date(val).strftime("%B %d, %Y")

def fmt_duration(seconds):
    if not seconds:
        return ""
    m = int(seconds) // 60
    return f"{m} min"

# ── Filename sanitiser ───────────────────────────────────────────────────────
def safe_name(s):
    for ch in r'<>:"/\|?*':
        s = s.replace(ch, "")
    return s.strip()

# ── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx(t, path):
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    title    = t["title"]
    d_prefix = date_prefix(t["date"])
    d_long   = date_long(t["date"])
    dur      = fmt_duration(t.get("duration"))

    # ── Meeting title ──────────────────────────────────────────────────────
    h = doc.add_heading(title, level=1)
    h.runs[0].font.size  = Pt(16)
    h.runs[0].font.bold  = True
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # ── Date / duration line ───────────────────────────────────────────────
    meta = doc.add_paragraph()
    meta_text = d_long + (f"  ·  {dur}" if dur else "")
    run = meta.add_run(meta_text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    meta.paragraph_format.space_after = Pt(14)

    # ── Horizontal rule ────────────────────────────────────────────────────
    hr_para = doc.add_paragraph()
    hr_para.paragraph_format.space_before = Pt(0)
    hr_para.paragraph_format.space_after  = Pt(14)
    pPr = hr_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Transcript body ────────────────────────────────────────────────────
    sentences = t.get("sentences") or []

    if not sentences:
        doc.add_paragraph("No transcript available.")
    else:
        # Group consecutive sentences from the same speaker
        turns = []
        for s in sentences:
            spk  = (s.get("speaker_name") or "Unknown").strip()
            text = (s.get("text") or "").strip()
            if not text:
                continue
            if turns and turns[-1][0] == spk:
                turns[-1][1].append(text)
            else:
                turns.append([spk, [text]])

        for spk, texts in turns:
            # Speaker name (bold)
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(10)
            sp.paragraph_format.space_after  = Pt(2)
            sr = sp.add_run(spk)
            sr.bold = True
            sr.font.size = Pt(11)
            sr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

            # Their words
            body = doc.add_paragraph(" ".join(texts))
            body.paragraph_format.space_before = Pt(0)
            body.paragraph_format.space_after  = Pt(4)
            body.runs[0].font.size = Pt(11)

    doc.save(path)

# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    print("Connecting to Fireflies API …")
    try:
        transcripts = fetch_all_transcripts()
    except Exception as e:
        print(f"ERROR: {e}")
        sys.exit(1)

    print(f"Found {len(transcripts)} meeting(s) up to March 30, 2026\n")

    ok = 0
    for t in transcripts:
        try:
            prefix   = date_prefix(t["date"])
            filename = f"{prefix} {safe_name(t['title'])}.docx"
            out_path = os.path.join(OUTPUT_DIR, filename)

            print(f"  → {filename}")
            build_docx(t, out_path)
            ok += 1
        except Exception as e:
            print(f"  ERROR on '{t.get('title')}': {e}")

    print(f"\n✓ Done — {ok}/{len(transcripts)} file(s) saved to:\n  {OUTPUT_DIR}\n")

if __name__ == "__main__":
    main()
